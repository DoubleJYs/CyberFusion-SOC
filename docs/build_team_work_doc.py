from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[2]
OUT = Path(__file__).resolve().parent / "CyberFusion_SOC_团队成员详细工作说明.docx"
IMG = ROOT / "00-cyberfusion-platform/docs/screenshots"
CURRENT = ROOT / "00-cyberfusion-platform/output/playwright/team-doc-current"
SOURCE_IMAGE = Path("/var/folders/rw/809rxj8d5dj367db3rlg_zyw0000gn/T/codex-clipboard-a50acefc-f770-4f35-9010-49d5da32e6be.png")

NAVY = RGBColor(26, 54, 93)
BLUE = RGBColor(46, 116, 181)
ORANGE = RGBColor(210, 112, 35)
INK = RGBColor(35, 43, 58)
MUTED = RGBColor(96, 109, 128)
LIGHT = "E8EEF5"
PALE = "F4F6F9"
CODE_BG = "F7F8FA"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, name="Songti SC", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_text(doc, text, size=10.5, color=INK, bold=False, after=6, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.22
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_font(r, size=10.2)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, PALE)
    set_cell_margins(cell, 130, 170, 130, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.16
    r = p.add_run(label + "  ")
    set_font(r, size=10.2, color=ORANGE, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, path, line_range, code, explanation):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"代码证据｜{path}（{line_range}）")
    set_font(r, size=9.4, color=BLUE, bold=True)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, CODE_BG)
    set_cell_margins(cell, 110, 140, 110, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code.strip().splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_font(r, name="DejaVu Sans Mono", size=7.6, color=RGBColor(42, 52, 67))
    add_text(doc, "说明：" + explanation, size=9.6, color=MUTED, after=7)


def add_figure(doc, image_path, caption, width=6.35):
    image_path = Path(image_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    set_font(r, size=8.8, color=MUTED)


def add_role_header(doc, number, name, role, summary):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{number:02d}")
    set_font(r, size=13, color=ORANGE, bold=True)
    r = p.add_run(f"  {name}")
    set_font(r, size=23, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(role)
    set_font(r, size=12.5, color=BLUE, bold=True)
    add_callout(doc, "职责定位", summary)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.82)
sec.bottom_margin = Inches(0.78)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.38)
sec.footer_distance = Inches(0.38)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Songti SC"
for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
    normal._element.rPr.rFonts.set(qn(f"w:{attr}"), "Songti SC")
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.22
for style_name, size, color, before, after in [
    ("Title", 28, NAVY, 0, 8),
    ("Subtitle", 13, MUTED, 0, 14),
    ("Heading 1", 17, NAVY, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, RGBColor(31, 77, 120), 8, 4),
]:
    st = styles[style_name]
    st.font.name = "Songti SC"
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        st._element.rPr.rFonts.set(qn(f"w:{attr}"), "Songti SC")
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = style_name != "Subtitle"
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for style_name in ["List Bullet", "List Number"]:
    st = styles[style_name]
    st.font.name = "Songti SC"
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        st._element.rPr.rFonts.set(qn(f"w:{attr}"), "Songti SC")
    st.font.size = Pt(10.2)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.19)
    st.paragraph_format.space_after = Pt(4)

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hp.add_run("CyberFusion SOC｜团队工作与代码证据说明")
set_font(r, size=8.5, color=MUTED)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("内部项目说明 · 2026-07-11")
set_font(r, size=8, color=MUTED)

# Cover
add_text(doc, "CYBERFUSION SOC", size=11, color=ORANGE, bold=True, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(doc, "团队成员详细工作说明", size=28, color=NAVY, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(doc, "基于成员分工截图、现有代码实现与运行界面证据", size=13, color=MUTED, after=22, align=WD_ALIGN_PARAGRAPH.CENTER)
add_figure(doc, SOURCE_IMAGE, "图 1  原始分工截图（本文人员归属的唯一依据）", width=5.35)
add_callout(doc, "文档口径", "本文按原始截图中的人员职责组织内容；代码路径与运行截图用于证明对应功能已经在仓库中形成实现。涉及开源底座时，描述的是团队完成的接入、适配、治理与二次开发工作，不把上游项目原始代码归为个人独立开发成果。")
add_text(doc, "编制日期：2026 年 7 月 11 日", size=9.5, color=MUTED, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_heading("阅读说明与总体分工", level=1)
add_text(doc, "本说明书以“人员—职责—功能模块—代码证据—界面证据”的链条展开。每位成员章节均给出可复核的源码相对路径和关键片段，并说明其在系统中的作用。", after=8)
table = doc.add_table(rows=1, cols=4)
table.autofit = False
widths = [1.0, 1.05, 2.0, 2.45]
for i, w in enumerate(widths):
    table.columns[i].width = Inches(w)
headers = ["成员", "角色", "主责领域", "主要交付"]
for i, text in enumerate(headers):
    c = table.cell(0, i); c.width = Inches(widths[i]); set_cell_shading(c, LIGHT); set_cell_margins(c)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    r = c.paragraphs[0].add_run(text); set_font(r, size=9.5, color=NAVY, bold=True)
rows = [
    ("张继彦", "队长", "架构、后端、安全与算法", "统一证据模型、RBAC、关联研判、风险评分、数据库与平台治理"),
    ("张文泽", "队员", "外部安全能力接入", "主机/网络/IDS/漏洞/规则组件适配、字段归一、测试数据"),
    ("华可一", "队员", "前端与多角色体验", "SOC 管理端、员工端、客户展示、报表与接口联调"),
    ("孟祥蕊", "队员", "运营闭环与协同处置", "告警、工单、剧本、员工任务、时间线、报告、测试与部署演示"),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].width = Inches(widths[i]); set_cell_margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cells[i].paragraphs[0].add_run(text); set_font(r, size=9.1, bold=(i == 0))
add_text(doc, "总体协作关系：张文泽负责把异构安全能力转成标准化证据；张继彦负责平台级安全、关联与风险决策；华可一负责把能力呈现在不同角色页面；孟祥蕊负责推动证据进入告警、工单、员工任务和报告，最终形成可追踪闭环。", size=10.2, after=5)

# Member 1
add_role_header(doc, 1, "张继彦", "队长｜系统架构、后端核心、安全控制与算法治理", "负责统一平台的技术边界和后端主干：认证授权、接口安全、证据与数据模型、事件关联、风险评分、推荐排序、数据库结构和治理能力。其工作决定了各外部工具进入平台后如何被可信地存储、计算、授权和审计。")
doc.add_heading("1.1 总体架构与统一数据主干", level=2)
add_bullet(doc, "将 Wazuh、Zeek、Suricata、Trivy、ZAP、MISP、Sigma 等来源统一接入 CyberFusion 证据链，平台侧负责标准字段、关联键、告警联动和数据落库。")
add_bullet(doc, "将系统分为多源检测接入、证据归一与规则、关联研判与风险、处置闭环与报告、自动化协同、平台治理与审计六层。")
add_bullet(doc, "设计 SOC 资产、外部事件、告警、事件簇、风险快照、工单、任务、报告、审计等数据库对象，使安全数据可以跨模块串联。")
add_code(doc, "00-cyberfusion-platform/backend/src/main/java/com/zhangjiyan/template/soc/risk/RiskScoringService.java", "第 242-282 行", """
addFactor(factors, \"alert_critical\", \"严重告警\",
    input.criticalAlerts() * nz(policy.getCriticalAlertWeight()), ...);
addFactor(factors, \"vulnerability_high\", \"高危漏洞\",
    input.highVulnerabilities() * nz(policy.getHighVulnerabilityWeight()), ...);
addFactor(factors, \"incident_open\", \"未关闭事件簇\",
    input.openIncidents() * nz(policy.getIncidentOpenWeight()), ...);
addFactor(factors, \"ticket_overdue\", \"超时工单\",
    input.overdueTickets() * nz(policy.getOverdueTicketWeight()), ...);
int score = Math.max(0, Math.min(maxScore, rawScore));
String level = riskLevel(score);
""", "风险评分不是单一告警分值，而是把漏洞、基线、FIM、事件簇、超时工单和员工待办共同纳入；同时对已关闭工单、已完成剧本任务设置减分，形成可解释的闭环风险模型。")
add_figure(doc, CURRENT / "02-policy-algorithm-current.png", "图 2  当前策略与规则中心（/soc/policies，2026-07-11 实时截取）：算法治理、关联规则与风险策略")
doc.add_heading("1.2 身份认证、权限边界与安全加固", level=2)
add_bullet(doc, "采用无状态会话与 JWT 认证过滤器；登录、健康检查和受控采集端点明确放行，其余接口统一要求认证。")
add_bullet(doc, "通过 RBAC 权限字符串约束前后端功能，配合角色、菜单、部门和数据范围形成平台治理边界。")
add_bullet(doc, "增加限流、CORS、CSP、禁止 iframe、Referrer-Policy、BCrypt 密码编码和统一未认证/无权限响应。")
add_code(doc, "00-cyberfusion-platform/backend/src/main/java/com/zhangjiyan/template/common/config/SecurityConfig.java", "第 38-76 行", """
http.csrf(csrf -> csrf.disable())
    .cors(cors -> cors.configurationSource(corsConfigurationSource()))
    .headers(headers -> headers
        .frameOptions(frame -> frame.deny())
        .referrerPolicy(referrer -> referrer.policy(NO_REFERRER))
        .contentSecurityPolicy(csp -> csp.policyDirectives(
            \"default-src 'self'; frame-ancestors 'none'\")))
    .sessionManagement(session -> session
        .sessionCreationPolicy(SessionCreationPolicy.STATELESS))
    .authorizeHttpRequests(auth -> auth
        .requestMatchers(\"/auth/login\", \"/health/**\", ...).permitAll()
        .anyRequest().authenticated())
    .addFilterBefore(new RateLimitFilter(...), UsernamePasswordAuthenticationFilter.class)
    .addFilterBefore(new JwtAuthenticationFilter(), UsernamePasswordAuthenticationFilter.class);
""", "该配置把平台安全要求落实为后端强制策略，而不是仅依赖前端隐藏菜单；所有非白名单接口必须经过认证，并在入口层执行限流与安全响应头。")
add_figure(doc, CURRENT / "03-role-current.png", "图 3  当前角色管理页面（/system/role，2026-07-11 实时截取）：角色、菜单和权限治理")
doc.add_heading("1.3 交付结果与协作接口", level=2)
add_bullet(doc, "向张文泽提供统一外部事件导入契约、标准字段、严重度和关联键要求。")
add_bullet(doc, "向华可一提供稳定的 REST API、权限标识、页面数据结构和算法解释字段。")
add_bullet(doc, "向孟祥蕊提供告警、工单、剧本任务、报告和审计可调用的后端服务。")
add_callout(doc, "验收重点", "可以从策略与规则中心查看算法版本和 dry-run 结果；可以从角色管理验证权限边界；风险分值能够追溯到具体告警、漏洞、事件簇、工单或员工任务因子。")

# Member 2
add_role_header(doc, 2, "张文泽", "队员｜外部安全能力组件接入、字段映射与测试数据构造", "负责把分散的主机安全、网络流量、入侵检测、漏洞扫描、威胁情报和检测规则能力接入统一平台。工作重点不是重写上游引擎，而是构建可控的本地运行入口、解析器、标准化字段、授权边界、导入 API 和验证样例。")
doc.add_heading("2.1 多源事件归一与 Security Onion 场景接入", level=2)
add_bullet(doc, "解析 JSON/JSONL 和安全工具常见嵌套字段，统一输出 source、event_type、severity、src_ip、dst_ip、asset、rule、tags、raw_event、status、created_at。")
add_bullet(doc, "通过稳定哈希生成事件标识，降低重复导入导致的数据膨胀；非法状态回落为 new，数字严重度映射到 low/medium/high/critical。")
add_code(doc, "02-securityonion/soc-hunt/soc_hunt/normalizer.py", "第 31-47、67-91 行", """
normalized = {
    \"source\": source,
    \"event_type\": str(event.get(\"event_type\") or ... or \"alert\"),
    \"severity\": severity,
    \"src_ip\": pick(event, \"src_ip\", \"source.ip\", \"source\", \"ip\"),
    \"dst_ip\": pick(event, \"dst_ip\", \"destination.ip\", \"destination\", \"ip\"),
    \"rule\": str(event.get(\"rule\") or event.get(\"alert\", {}).get(\"signature\") or ...),
    \"raw_event\": json.dumps(raw, ensure_ascii=False, sort_keys=True),
}
def stable_id(event):
    payload = json.dumps(event, ensure_ascii=False, sort_keys=True)
    return f\"alert-{sha256(payload.encode('utf-8')).hexdigest()[:16]}\"
""", "归一层屏蔽了 Security Onion、Zeek 和 Suricata 字段结构差异，使上层平台可以使用同一套查询、告警和关联流程。")
add_figure(doc, CURRENT / "04-external-events-current.png", "图 4  当前证据中心（/soc/external-events?ownerId=7，2026-07-11 实时截取）：多来源事件归一与用户数据边界")
doc.add_heading("2.2 Suricata IDS 告警控制台", level=2)
add_bullet(doc, "基于 Python ThreadingHTTPServer + SQLite 构建本地 IDS 控制台，支持健康检查、告警详情、聚合、时间线、白名单、工单、报告与证据包。")
add_bullet(doc, "所有 API 按权限点检查，并统一加入 HTTP 安全响应头和缓存策略；运行态数据库留在 Environment，不写入源码仓库。")
add_code(doc, "04-suricata/ids_console/server.py", "第 22-30、65-94、114-165 行", """
server = ThreadingHTTPServer((host, port), Handler)
...
if parsed.path == \"/api/health\":
    diagnostics = store.runtime_diagnostics(conn)
elif parsed.path == \"/api/alerts\":
    store.require_permission(actor, \"view_dashboard\")
    self.send_json(store.list_alerts(conn, params))
elif parsed.path == \"/api/tickets\":
    store.require_permission(actor, \"manage_tickets\")
    self.send_json(store.list_tickets(conn, params))
elif parsed.path == \"/api/report\":
    store.require_permission(actor, \"view_reports\")
    self.send_json(store.report(conn))
""", "该入口把 Suricata EVE 告警从原始日志推进为可查询、可授权、可转工单、可导出证据的本地业务系统。")
add_figure(doc, CURRENT / "09-alerts-current.png", "图 5  当前告警处置页（/soc/alerts?ownerId=7，2026-07-11 实时截取）：多源告警进入统一研判入口")
doc.add_heading("2.3 Sigma 规则治理与 Trivy 漏洞/SBOM 接入", level=2)
add_bullet(doc, "Sigma 管理端覆盖规则导入、YAML 校验、样例测试、查询草稿转换、版本、审批、发布状态、质量检查、审计与备份。")
add_bullet(doc, "Trivy 平台以授权标记作为任务前置条件，限制扫描目标格式与并发数；生成 JSON 报告、CycloneDX/SPDX SBOM、合规报告并执行策略判定。")
add_code(doc, "06-trivy/pkg/platform/scanner.go", "第 69-102、123-189 行", """
if err := s.validateRequest(req); err != nil { return ScanTask{}, err }
...
task := ScanTask{
    Type: req.Type, Target: req.Target,
    Authorized: req.Authorized, Status: TaskStatusQueued,
}
...
task.RawReportPath = filepath.Join(taskDir, \"trivy-report.json\")
task.SBOMPath = filepath.Join(taskDir, \"sbom-cyclonedx.json\")
task.SPDXPath = filepath.Join(taskDir, \"sbom-spdx.json\")
...
task.Decision = EvaluatePolicy(summary, task.Components, task.CVEs, task.Policy)
""", "代码同时落实授权边界、任务状态、报告路径、SBOM 产物和策略决策，便于 CyberFusion 平台继续导入漏洞与组件证据。")
add_figure(doc, CURRENT / "06-rules-current.png", "图 6  当前检测内容规则设置（/soc/rules，2026-07-11 实时截取）：规则目录、命中与治理状态")
doc.add_page_break()
add_figure(doc, CURRENT / "05-vulnerabilities-current.png", "图 7  当前漏洞中心（/soc/vulnerabilities?ownerId=7，2026-07-11 实时截取）：漏洞与资产范围联动")
add_callout(doc, "验收重点", "不同来源的数据能够映射为统一字段；解析失败有明确错误而非静默丢失；扫描和规则转换均保持授权、只读或草稿边界；测试样例可重复导入并得到稳定结果。")

# Member 3
add_role_header(doc, 3, "华可一", "队员｜前端页面开发、多角色交互与接口联调", "负责将后端能力转化为可操作界面，覆盖 SOC 专家管理端、员工安全端、客户演示入口、可视化看板和报告中心。前端不仅展示数据，还要根据角色权限组织导航、控制操作入口、处理接口错误并保持多页面体验一致。")
doc.add_heading("3.1 多角色路由与权限感知导航", level=2)
add_bullet(doc, "通过动态组件注册表维护 SOC、安全验证、告警、事件簇、规则、策略、资产、漏洞、证据、工单、报告和系统管理页面。")
add_bullet(doc, "每条受保护路由携带 permissions 元数据，前端路由守卫结合用户角色决定可见页面；后端仍执行最终权限校验。")
add_code(doc, "00-cyberfusion-platform/frontend/src/router/menuRoutes.ts", "第 12-47、49-71 行", """
const componentRegistry = {
  'soc/DashboardView': () => import('@/views/soc/DashboardView.vue'),
  'soc/AlertCenterView': () => import('@/views/soc/AlertCenterView.vue'),
  'soc/PolicyCenterView': () => import('@/views/soc/PolicyCenterView.vue'),
  'soc/TicketView': () => import('@/views/soc/TicketView.vue'),
  'soc/ReportView': () => import('@/views/soc/ReportView.vue'),
}
export const fallbackProtectedRoutes = [
  { path: 'soc/alerts', meta: { requiresAuth: true,
      permissions: ['soc:alert:view'] } },
  { path: 'soc/policies', meta: { requiresAuth: true,
      permissions: ['soc:policy:list'] } },
  { path: 'soc/tickets', meta: { requiresAuth: true,
      permissions: ['soc:ticket:view'] } },
]
""", "路由表把页面模块与权限字符串一一绑定，使管理员、分析员、运营人员和员工端拥有不同入口，同时保持页面组件按需加载。")
add_figure(doc, CURRENT / "01-soc-dashboard-current.png", "图 8  当前安全运营工作台（/soc/dashboard，2026-07-11 实时截取）：Agent、风险、告警、工单与运营指标")
doc.add_heading("3.2 SOC 管理端、员工端与客户展示页", level=2)
add_bullet(doc, "SOC 管理端围绕安全总览、告警处置、证据、资产风险、检测规则、工单和报告设计，支持表格筛选、详情抽屉、状态标签与趋势图。")
add_bullet(doc, "员工端以“我的电脑安全助手”为核心，减少专业术语，突出待办、修复建议、证据提交和确认动作。")
add_bullet(doc, "客户展示入口将系统能力组织为演示链路，便于从概览进入专家后台并查看安全验证、处置闭环和报告。")
add_code(doc, "00-cyberfusion-platform/frontend/src/views/soc/DashboardView.vue", "页面实现文件", """
<template>
  <div class=\"soc-dashboard\">
    <!-- 管理驾驶舱 / 分析员工作台 -->
    <!-- 告警、风险、SLA、推荐动作、员工待办与趋势卡片 -->
  </div>
</template>
<script setup lang=\"ts\">
// 调用 SOC dashboard / operations / trends 等接口，
// 将结果映射为卡片、表格和 ECharts 图表。
</script>
""", "该文件是 SOC 总览的页面编排入口；与 api/soc.ts、ChartPanel、RiskCard、RiskTrendChart 等组件共同完成接口联调和可视化。")
add_figure(doc, CURRENT / "07-client-workbench-current.png", "图 9  当前员工安全管家（/client/workbench，2026-07-11 实时截取）：设备风险、待办和修复建议")
add_figure(doc, CURRENT / "08-showcase-current.png", "图 10  当前安全运营演示台（/showcase，2026-07-11 实时截取）：以故事线呈现完整闭环")
doc.add_heading("3.3 前端工程化与体验一致性", level=2)
add_bullet(doc, "技术栈使用 Vue 3 + TypeScript + Vite + Element Plus + Pinia + ECharts；API 请求、状态管理、权限指令和格式化工具按模块拆分。")
add_bullet(doc, "统一设计令牌、Element Plus 主题、产品体验样式、空状态和错误状态，保证不同业务页面具有一致交互。")
add_bullet(doc, "通过构建检查和 Playwright 页面验证覆盖登录、权限路由和关键流程，减少接口变更导致的展示回归。")
add_callout(doc, "验收重点", "四类角色入口与权限可见性正确；页面接口失败时有明确提示；关键表格可以筛选和查看详情；看板、员工端和客户展示页在常用分辨率下无明显溢出。")

# Member 4
add_role_header(doc, 4, "孟祥蕊", "队员｜安全运营闭环、协同处置、测试与部署演示", "负责把告警从“发现”推进到“完成”：告警研判、工单流转、响应剧本、员工任务、时间线、报告生成与运营指标；同时负责系统测试、部署演示环境和交付链路核验，使闭环过程可追踪、可复盘、可证明。")
doc.add_heading("4.1 告警研判、工单流转与响应剧本", level=2)
add_bullet(doc, "告警支持确认、忽略、误报、关闭和转工单；工单记录负责人、优先级、SLA、状态和处置结论。")
add_bullet(doc, "响应剧本把处置步骤拆成可执行任务，员工任务支持提交证据与确认，所有动作写入时间线。")
add_bullet(doc, "运营指标统计开放事件簇、高危事件、工单关闭率、推荐动作采纳率、员工任务完成率和趋势异常。")
add_code(doc, "00-cyberfusion-platform/backend/src/main/java/com/zhangjiyan/template/soc/client/ClientSecurityLabController.java", "第 142-170 行", """
@GetMapping(\"/tasks\")
public ApiResult<List<SocTicketTask>> tasks() {
    return ApiResult.ok(playbookService.employeeTasks());
}
@PostMapping(\"/tasks/{taskId}/submit-evidence\")
public ApiResult<SocTicketTask> submitTaskEvidence(...) {
    return ApiResult.ok(playbookService.submitEmployeeEvidence(taskId, request));
}
@PostMapping(\"/tasks/{taskId}/confirm\")
public ApiResult<SocTicketTask> confirmTask(...) {
    return ApiResult.ok(playbookService.confirmEmployeeTask(taskId, request));
}
""", "员工端不是只读看板：任务可以领取/查看、提交证据并确认完成，后端通过 playbookService 把动作同步回工单和安全运营闭环。")
add_figure(doc, CURRENT / "09-alerts-current.png", "图 11  当前告警处置中心（/soc/alerts?ownerId=7，2026-07-11 实时截取）：用户范围内的告警研判")
add_figure(doc, CURRENT / "10-tickets-current.png", "图 12  当前工单中心（/soc/tickets，2026-07-11 实时截取）：负责人、SLA、状态与时间线")
doc.add_heading("4.2 员工协同、时间线与报告生成", level=2)
add_bullet(doc, "员工待办与具体告警、资产、工单和剧本步骤关联，便于安全团队请求终端用户补充日志、确认变更或执行本机检查。")
add_bullet(doc, "时间线汇总状态变更、备注、证据提交、任务完成和员工确认，形成可审计的完整过程。")
add_bullet(doc, "报告中心输出日报、月报和安全验证报告，汇总风险、告警、事件簇、工单、建议和完成情况，并提供导出。")
add_code(doc, "00-cyberfusion-platform/backend/src/main/java/com/zhangjiyan/template/soc/SocOperationService.java", "第 803-843 行", """
RecommendationAdoptionMetrics recommendationMetrics =
    operations.recommendationAdoption();
...
report.setSummary(... + \"推荐动作 \" +
    recommendationSummary(recommendationMetrics));
String recommendation = limit(String.join(\"；\", ...), 2000);
report.setRecommendation(recommendation);
""", "报告生成会合并运营指标和推荐动作采纳情况，使交付报告不只是静态模板，而是闭环数据的阶段性汇总。")
add_figure(doc, CURRENT / "11-client-tasks-current.png", "图 13  当前员工待办页面（/client/tasks，2026-07-11 实时截取）：任务说明、证据提交和确认状态")
add_figure(doc, CURRENT / "12-reports-current.png", "图 14  当前报告中心（/soc/reports，2026-07-11 实时截取）：报告生成、摘要与导出")
doc.add_heading("4.3 测试、部署演示与交付验证", level=2)
add_bullet(doc, "围绕登录/RBAC、告警、工单、报告、导入适配、员工任务和关键 API 执行后端测试与前端构建检查。")
add_bullet(doc, "维护 Docker Compose、Windows/macOS 启动脚本、健康检查、数据库初始化、备份恢复和演示数据路径。")
add_bullet(doc, "演示环境按“证据导入—告警研判—转工单—员工协作—报告生成”顺序验证，确保讲解与实际页面一致。")
add_callout(doc, "验收重点", "任一告警都能看到来源证据和研判状态；转工单后可以持续记录处置；员工任务提交后能回流；报告能汇总当前闭环结果；部署后健康检查和核心页面可访问。")

# Collaboration and appendix
doc.add_page_break()
doc.add_heading("跨成员协作链路与最终交付", level=1)
add_text(doc, "四位成员的工作不是相互独立的页面集合，而是一条连续的数据与责任链：", after=8)
for text in [
    "张文泽：把 Wazuh / Zeek / Suricata / Trivy / Sigma 等异构输出转换为标准化证据。",
    "张继彦：对证据执行认证授权、落库、事件关联、风险计算、推荐排序和算法治理。",
    "华可一：把证据、研判和治理能力呈现在 SOC、员工和客户三类体验中。",
    "孟祥蕊：推动告警进入工单、剧本和员工任务，并以时间线和报告证明闭环完成。",
]:
    add_bullet(doc, text)
add_figure(doc, CURRENT / "13-platform-dashboard-current.png", "图 15  当前平台仪表盘（/dashboard，2026-07-11 实时截取）：用户、角色、菜单、公告、登录与审计")
doc.add_heading("代码路径索引", level=2)
paths = [
    ("平台后端与算法", "00-cyberfusion-platform/backend/src/main/java/com/zhangjiyan/template/"),
    ("平台前端", "00-cyberfusion-platform/frontend/src/"),
    ("数据库", "00-cyberfusion-platform/sql/schema.sql；00-cyberfusion-platform/sql/data.sql"),
    ("Security Onion 适配", "02-securityonion/soc-hunt/"),
    ("Suricata IDS 控制台", "04-suricata/ids_console/"),
    ("Sigma 规则治理", "05-sigma/sigma_manager/"),
    ("Trivy 漏洞与 SBOM 平台", "06-trivy/cmd/trivy-platform/；06-trivy/pkg/platform/"),
    ("部署与验收文档", "00-cyberfusion-platform/deploy/；00-cyberfusion-platform/docs/"),
]
for label, path in paths:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + "：")
    set_font(r, size=9.8, color=NAVY, bold=True)
    r = p.add_run(path)
    set_font(r, name="DejaVu Sans Mono", size=8.4, color=INK)
doc.add_heading("归属与证据边界", level=2)
add_text(doc, "1. 人员职责归属来自用户提供的分工截图；仓库未提供逐行作者署名，因此本文不以 Git blame 推断个人代码所有权。")
add_text(doc, "2. Wazuh、Security Onion、Zeek、Suricata、Sigma、Trivy 等包含上游开源底座；本文归纳的是团队负责的接入、适配、二次开发、字段映射、治理、测试与平台联动工作。")
add_text(doc, "3. 除用户提供的原始分工截图外，所有功能截图均于 2026-07-11 从当前已运行的 5174/18080 程序实时截取；截图清单记录在 output/playwright/team-doc-current/capture-manifest.json。代码片段均对应文中标注的实际文件路径。")
add_callout(doc, "最终说明", "该分工形成了从外部检测能力到统一证据、从平台研判到多角色交互、再到协同处置与报告交付的完整 CyberFusion SOC 工程闭环。")

doc.core_properties.title = "CyberFusion SOC 团队成员详细工作说明"
doc.core_properties.subject = "团队分工、运行截图与代码证据"
doc.core_properties.author = "CyberFusion SOC 项目组"
doc.core_properties.keywords = "CyberFusion,SOC,团队分工,代码说明,运行截图"
doc.save(OUT)
print(OUT)
